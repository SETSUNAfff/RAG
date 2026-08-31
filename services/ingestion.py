"""文档解析与知识库入库服务。"""

from __future__ import annotations

import asyncio
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from crud.mysql.chunks import create_chunks, hard_delete_chunks_by_document
from models.mysql import Document
from schemas.mysql import ChunkCreate, DocumentStatus
from services.cache import invalidate_search_cache
from services.embeddings import get_embedding_model
from services.locks import ingestion_lock


SUPPORTED_FILE_TYPES = ["pdf", "docx", "md", "markdown", "html", "htm", "txt"]
SUPPORTED_EXTENSIONS = {f".{extension}" for extension in SUPPORTED_FILE_TYPES}
_TEXT_ENCODINGS = ("utf-8-sig", "gb18030", "utf-16", "latin-1")


# 文档上传，分析格式，清洗数据
def _decode_text(data: bytes) -> str:
    for encoding in _TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_text(file_name: str, data: bytes) -> str:
    """根据文件扩展名从上传文件字节流中提取可检索文本。"""
    extension = Path(file_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型：{extension or '(无扩展名)'}，"
            f"支持：{', '.join(SUPPORTED_FILE_TYPES)}"
        )

    if extension == ".pdf":
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages)
        if not text.strip():
            raise ValueError("PDF 未提取到文本，请确认文件不是扫描件")
    elif extension == ".docx":
        document = Document(BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    elif extension in {".html", ".htm"}:
        soup = BeautifulSoup(data, "lxml")
        text = soup.get_text("\n", strip=True)
    else:
        text = _decode_text(data)

    return _normalize_text(text)


# 分割文本
def split_text(text: str, chunk_size: int = 500, chunk_overlap: int = 80) -> list[str]:
    """把长文本按语义边界切分，供 embedding 和 Milvus 入库使用。"""
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "?", " ", ""],
    )
    return splitter.split_text(text)


# 上传文件入库：MySQL chunks 先拿到自增 ID，Milvus 再使用同一 ID。
async def ingest_uploaded_file(
    db: AsyncSession,
    file_name: str,
    data: bytes,
    document_id: int,
    *,
    replace_existing: bool = False,
    page_no: int = 0,
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """解析上传文件并同时写入 MySQL 与 Milvus，返回向量入库结果。"""
    from crud.milvus.knowledge_chunks import (
        delete_document_chunks,
        insert_knowledge_chunks,
    )

    document = await db.get(Document, document_id)
    if document is None:
        raise ValueError("Document not found")

    embedding_model = await asyncio.to_thread(get_embedding_model)

    text = extract_text(file_name, data)
    chunks = split_text(text)
    if not chunks:
        raise ValueError("文件未提取到有效文本")

    embeddings = await asyncio.to_thread(embedding_model.encode, chunks)
    extension = Path(file_name).suffix.lower().lstrip(".")
    page_nos = [page_no] * len(chunks)
    chunk_metadata = [
        {**(metadata or {}), "file_name": file_name, "source_type": extension}
        for _ in chunks
    ]
    chunk_creates = [
        ChunkCreate(
            document_id=document_id,
            content=chunk,
            page_no=page_no,
            token_count=max(1, len(chunk.split())),
            metadata=chunk_metadata[index],
        )
        for index, chunk in enumerate(chunks)
    ]

    async with ingestion_lock:
        try:
            if replace_existing:
                await hard_delete_chunks_by_document(db, document_id)
                await asyncio.to_thread(delete_document_chunks, document_id)

            db_chunks = await create_chunks(db, chunk_creates)
            chunk_ids = [chunk.id for chunk in db_chunks]

            result = await asyncio.to_thread(
                insert_knowledge_chunks,
                chunk_id=chunk_ids,
                document_id=document_id,
                embeddings=embeddings.tolist(),
                content_text=chunks,
                page_no=page_nos,
                metadata=chunk_metadata,
            )

            document.status = DocumentStatus.READY.value
            if replace_existing:
                document.version += 1
            document.error_message = None
            await db.commit()
            # 提交后刷新服务端生成的 updated_at，避免响应序列化时触发异步懒加载。
            await db.refresh(document)
            # 文档内容变更后清空 Redis 检索缓存，避免命中旧结果。
            await invalidate_search_cache()
            return result
        except Exception as exc:
            await db.rollback()
            try:
                await asyncio.to_thread(delete_document_chunks, document_id)
            except Exception:
                pass
            document = await db.get(Document, document_id)
            if document is not None:
                document.status = DocumentStatus.FAILED.value
                document.error_message = str(exc)
                await db.commit()
                await db.refresh(document)
            raise
